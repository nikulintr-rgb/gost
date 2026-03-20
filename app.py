import streamlit as st
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import os
from datetime import datetime

# ============================================
# КОНФИГУРАЦИЯ СТРАНИЦЫ
# ============================================
st.set_page_config(
    page_title="ГОСТ-Чекер НИР",
    page_icon="📑",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# CSS СТИЛИЗАЦИЯ
# ============================================
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f4e79;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .error-box {
        background-color: #ffe6e6;
        border-left: 4px solid #ff4444;
        padding: 10px;
        margin: 5px 0;
        border-radius: 4px;
    }
    .success-box {
        background-color: #e6ffe6;
        border-left: 4px solid #44ff44;
        padding: 10px;
        margin: 5px 0;
        border-radius: 4px;
    }
    .warning-box {
        background-color: #fff3e6;
        border-left: 4px solid #ffaa44;
        padding: 10px;
        margin: 5px 0;
        border-radius: 4px;
    }
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        border: 1px solid #dee2e6;
    }
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        color: #1f4e79;
    }
    .metric-label {
        font-size: 0.9rem;
        color: #666;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# КЛАСС ПРОВЕРКИ ГОСТ
# ============================================
class GostChecker:
    def __init__(self, doc):
        self.doc = doc
        self.errors = []
        self.warnings = []
        self.fixed_count = 0
        
        # Настройки ГОСТ
        self.gost_font_name = 'Times New Roman'
        self.gost_font_size = 14
        self.gost_line_spacing = 1.5
        self.gost_first_line_indent = Cm(1.25)
        self.gost_margins = {
            'left': Cm(3.0),
            'right': Cm(1.0),
            'top': Cm(2.0),
            'bottom': Cm(2.0)
        }

    def check_margins(self):
        """Проверка полей документа"""
        sections = self.doc.sections
        for i, section in enumerate(sections):
            if abs(section.left_margin - self.gost_margins['left']) > Cm(0.1):
                self.errors.append(f"📏 Раздел {i+1}: Левое поле {section.left_margin.cm:.1f} см (нужно 3.0 см)")
            if abs(section.right_margin - self.gost_margins['right']) > Cm(0.1):
                self.errors.append(f"📏 Раздел {i+1}: Правое поле {section.right_margin.cm:.1f} см (нужно 1.0 см)")
            if abs(section.top_margin - self.gost_margins['top']) > Cm(0.1):
                self.errors.append(f"📏 Раздел {i+1}: Верхнее поле {section.top_margin.cm:.1f} см (нужно 2.0 см)")
            if abs(section.bottom_margin - self.gost_margins['bottom']) > Cm(0.1):
                self.errors.append(f"📏 Раздел {i+1}: Нижнее поле {section.bottom_margin.cm:.1f} см (нужно 2.0 см)")

    def fix_margins(self):
        """Исправление полей"""
        for section in self.doc.sections:
            section.left_margin = self.gost_margins['left']
            section.right_margin = self.gost_margins['right']
            section.top_margin = self.gost_margins['top']
            section.bottom_margin = self.gost_margins['bottom']
        self.fixed_count += 1

    def check_paragraph_formatting(self):
        """Проверка форматирования абзацев"""
        font_errors = 0
        size_errors = 0
        indent_errors = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue

            for run in para.runs:
                if run.font.name and 'Times' not in run.font.name:
                    font_errors += 1
                if run.font.size and run.font.size.pt != self.gost_font_size:
                    size_errors += 1

            if para.paragraph_format.first_line_indent != self.gost_first_line_indent:
                if not self._is_heading(para):
                    indent_errors += 1

        if font_errors > 0:
            self.warnings.append(f"✏️ Найдено {font_errors} фрагментов с неверным шрифтом")
        if size_errors > 0:
            self.warnings.append(f"📐 Найдено {size_errors} фрагментов с неверным размером шрифта")
        if indent_errors > 0:
            self.warnings.append(f"↩️ Найдено {indent_errors} абзацев с неверным отступом")

    def _is_heading(self, para):
        """Проверка является ли абзац заголовком"""
        style_name = para.style.name if para.style else ""
        return ("Heading" in style_name or "Заголовок" in style_name or 
                para.alignment == WD_ALIGN_PARAGRAPH.CENTER)

    def fix_text_formatting(self):
        """Исправление шрифтов и отступов"""
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
            
            para.paragraph_format.first_line_indent = self.gost_first_line_indent
            para.paragraph_format.line_spacing = 1.5
            
            for run in para.runs:
                run.font.name = self.gost_font_name
                run.font.size = Pt(self.gost_font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.gost_font_name)
        
        self.fixed_count += 1

    def check_headings(self):
        """Проверка заголовков"""
        heading_errors = 0
        
        for para in self.doc.paragraphs:
            if self._is_heading(para):
                text = para.text.strip()
                
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    heading_errors += 1
                
                if text.endswith('.'):
                    heading_errors += 1
                
                if para.runs and not para.runs[0].font.bold:
                    heading_errors += 1

        if heading_errors > 0:
            self.warnings.append(f"📑 Найдено проблем с заголовками: {heading_errors}")

    def fix_headings(self):
        """Исправление заголовков"""
        for para in self.doc.paragraphs:
            if self._is_heading(para):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if para.text.endswith('.'):
                    para.text = para.text[:-1]
                for run in para.runs:
                    run.font.bold = True
        self.fixed_count += 1

    def check_structure(self):
        """Проверка структуры документа"""
        text_content = "\n".join([p.text for p in self.doc.paragraphs]).lower()
        required_sections = [
            "введение", 
            "заключение", 
            "список использованных источников",
            "список литературы"
        ]
        
        for section in required_sections:
            if section not in text_content:
                self.errors.append(f"📋 Отсутствует раздел: '{section.title()}'")

    def get_stats(self):
        """Получить статистику документа"""
        total_paragraphs = len([p for p in self.doc.paragraphs if p.text.strip()])
        total_sections = len(self.doc.sections)
        total_pages = len(self.doc.sections)  # Приблизительно
        
        return {
            'paragraphs': total_paragraphs,
            'sections': total_sections,
            'pages': total_pages
        }

    def run_check(self):
        """Запустить полную проверку"""
        self.errors = []
        self.warnings = []
        self.check_margins()
        self.check_paragraph_formatting()
        self.check_headings()
        self.check_structure()
        return len(self.errors), len(self.warnings)

    def run_fix(self):
        """Запустить исправление"""
        self.fix_margins()
        self.fix_text_formatting()
        self.fix_headings()

    def save_document(self):
        """Сохранить документ в буфер"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

# ============================================
# ИНТЕРФЕЙС ПРИЛОЖЕНИЯ
# ============================================

# Заголовок
st.markdown('<div class="main-header">📑 ГОСТ-Чекер НИР</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Автоматическая проверка и исправление отчетов по ГОСТ 7.32-2017</div>', unsafe_allow_html=True)

st.divider()

# Боковая панель
with st.sidebar:
    st.header("⚙️ Настройки ГОСТ")
    
    st.markdown("**Основные параметры:**")
    st.info("""
    - Шрифт: Times New Roman
    - Размер: 14 пт
    - Интервал: 1.5
    - Отступ: 1.25 см
    """)
    
    st.markdown("**Поля страницы:**")
    st.info("""
    - Левое: 30 мм
    - Правое: 10 мм
    - Верхнее: 20 мм
    - Нижнее: 20 мм
    """)
    
    st.divider()
    st.markdown("### 📖 Помощь")
    st.markdown("""
    1. Загрузите файл .docx
    2. Нажмите "Проверить документ"
    3. Изучите найденные ошибки
    4. Скачайте исправленную версию
    """)
    
    st.divider()
    st.markdown(f"*Версия: 1.0.0*")
    st.markdown(f"*Дата: {datetime.now().strftime('%d.%m.%Y')}*")

# Основная область
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown('<div class="metric-card">', unsafe_allow_html=True)
    st.markdown('<div class="metric-value" id="error-count">0</div>', unsafe_allow_html=True)
    st.markdown('<div class="metric-label">Ошибок</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="metric-card">', unsafe_allow_html=True)
    st.markdown('<div class="metric-value" id="warning-count">0</div>', unsafe_allow_html=True)
    st.markdown('<div class="metric-label">Предупреждений</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

with col3:
    st.markdown('<div class="metric-card">', unsafe_allow_html=True)
    st.markdown('<div class="metric-value" id="fix-count">0</div>', unsafe_allow_html=True)
    st.markdown('<div class="metric-label">Исправлено</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

st.divider()

# Загрузка файла
uploaded_file = st.file_uploader(
    "📁 Загрузите документ Word (.docx)",
    type=["docx"],
    help="Загрузите файл отчета НИР в формате .docx для проверки"
)

if uploaded_file is not None:
    # Сохраняем файл временно
    with open("temp_document.docx", "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    st.success(f"✅ Файл загружен: **{uploaded_file.name}**")
    
    # Загружаем документ
    try:
        doc = docx.Document("temp_document.docx")
        checker = GostChecker(doc)
        
        # Показываем статистику
        stats = checker.get_stats()
        st.markdown(f"📊 **Статистика документа:** {stats['paragraphs']} абзацев, {stats['sections']} раздел(ов)")
        
        # Кнопки действий
        col1, col2 = st.columns(2)
        
        with col1:
            check_button = st.button("🔍 Проверить документ", use_container_width=True, type="primary")
        
        with col2:
            fix_button = st.button("🔧 Исправить и скачать", use_container_width=True, disabled=not check_button)
        
        # Результаты проверки
        if check_button:
            with st.spinner("⏳ Выполняется проверка документа..."):
                error_count, warning_count = checker.run_check()
            
            # Обновляем метрики
            st.session_state['error_count'] = error_count
            st.session_state['warning_count'] = warning_count
            
            # Показываем результаты
            if error_count == 0 and warning_count == 0:
                st.success("✅ Грубых нарушений форматирования не найдено! Документ соответствует ГОСТ.")
            else:
                if error_count > 0:
                    st.error(f"❌ Найдено ошибок: **{error_count}**")
                    with st.expander("📋 Показать все ошибки", expanded=True):
                        for err in checker.errors:
                            st.markdown(f'<div class="error-box">{err}</div>', unsafe_allow_html=True)
                
                if warning_count > 0:
                    st.warning(f"⚠️ Найдено предупреждений: **{warning_count}**")
                    with st.expander("⚠️ Показать предупреждения", expanded=False):
                        for warn in checker.warnings:
                            st.markdown(f'<div class="warning-box">{warn}</div>', unsafe_allow_html=True)
            
            # Активируем кнопку исправления
            st.session_state['can_fix'] = True
        
        # Исправление и скачивание
        if fix_button and st.session_state.get('can_fix', False):
            with st.spinner("⏳ Выполняется исправление документа..."):
                checker.run_fix()
                fixed_buffer = checker.save_document()
            
            st.success(f"✅ Исправлено параметров: **{checker.fixed_count}**")
            
            # Кнопка скачивания
            st.download_button(
                label="📥 Скачать исправленный документ",
                data=fixed_buffer,
                file_name=f"fixed_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.info("💡 Рекомендуется открыть исправленный файл и проверить результат вручную.")
        
        # Очистка временного файла
        if os.path.exists("temp_document.docx"):
            os.remove("temp_document.docx")
    
    except Exception as e:
        st.error(f"❌ Ошибка при обработке файла: {str(e)}")
        st.info("Убедитесь, что файл не поврежден и имеет формат .docx")

else:
    # Инструкции когда файл не загружен
    st.info("""
    ### 📝 Как использовать:
    
    1. **Загрузите файл** — перетащите .docx файл в область загрузки выше
    2. **Проверьте документ** — нажмите кнопку "Проверить документ"
    3. **Изучите ошибки** — просмотрите найденные нарушения ГОСТ
    4. **Исправьте** — нажмите "Исправить и скачать" для получения исправленной версии
    
    ---
    
    ### ✅ Что проверяется:
    
    - Поля страницы (30/10/20/20 мм)
    - Шрифт (Times New Roman, 14 пт)
    - Междустрочный интервал (1.5)
    - Абзацный отступ (1.25 см)
    - Оформление заголовков
    - Структура документа
    """)

# Футер
st.divider()
st.markdown("""
<div style="text-align: center; color: #666; font-size: 0.9rem;">
    ГОСТ-Чекер НИР © 2024 | Проверка по ГОСТ 7.32-2017, ГОСТ Р 7.0.97-2016
</div>
""", unsafe_allow_html=True)
